"""Editable DOCX rendering of the Teroxx allocation proposal.

DOCX is the primary working format for the Investment Advisory team:
the auto-rendered file is a starting point that advisors layer with
client-specific wishes, summaries, execution analysis and comments
before sending. Per Jannick Bröring (2026-05-14): "with a PDF we can
not work, the outcome should be a docx".

The renderer takes the same ``ctx`` dict that ``proposal.render_pdf``
consumes (built by :func:`build_proposal_context`). Exhibits authored
as SVG (donut, regime gauge, tier bar) are rasterised to PNG via
cairosvg and embedded as inline images so the advisor opens a Word
document that already looks done.

Brand fonts (Söhne, Sometimes Times) are declared in the document
defaults. Word will pick them up if the advisor has the fonts
installed locally; otherwise it falls back to Calibri / Cambria
gracefully — the document remains editable either way.
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, Iterable, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Mm, Pt, RGBColor

from app.pdf.i18n import (
    DEFAULT as I18N_DEFAULT,
    t,
    profile_label,
    tier_label,
    regime_label,
)


# ── Brand palette (Short Brand Guideline VS1.0 §3.2) ────────────────
NIGHTBLUE = RGBColor(0x01, 0x06, 0x26)
DEEP_INDIGO = RGBColor(0x06, 0x0D, 0x43)
ELECTRIC_SKY = RGBColor(0x0B, 0x68, 0x8C)
SANDSTONE = RGBColor(0xBF, 0xB3, 0xA8)
SUNSET_EMBER = RGBColor(0xD0, 0x66, 0x43)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Brand picks Deep Indigo over pure black for text (the brand "Color
# Story" pairs the Deep Indigo with Nightblue as the foundation).
TEXT_BODY = DEEP_INDIGO
TEXT_MUTED = RGBColor(0x6B, 0x6F, 0x82)  # desaturated Deep Indigo at ~55%

# Brand palette has no green; gains map to Electric Sky and losses to
# Sunset Ember (mirrors the same decision in palette.py / proposal.css).
GAIN = ELECTRIC_SKY
LOSS = SUNSET_EMBER

# Sandstone tints for backgrounds. Hex strings (not RGBColor) because
# _shade_cell takes raw hex.
SANDSTONE_50_HEX = "F4F1ED"     # zebra-stripe row
CREAM_HEX = "ECE8E5"            # disclaimer / appendix blocks


# ── Brand typography (Short Brand Guideline VS1.0 §5.1, §5.2) ───────
# Word looks up these by the font family name embedded in the OTF. The
# Söhne weights ship as separate families (Söhne / Söhne Leicht /
# Söhne Kräftig / Söhne Halbfett / Söhne Fett); set font.name to the
# weight you want and keep bold=False so Word doesn't try to synth.
FONT_HEADING = "Sometimes Times Medium"     # Big Headline, action titles
FONT_BODY = "Söhne"                          # Bodytext = Söhne Buch
FONT_LEICHT = "Söhne Leicht"                # Sublines
FONT_KRAFTIG = "Söhne Kräftig"              # Small special headings, kickers
FONT_HALBFETT = "Söhne Halbfett"            # Semi-bold accents
FONT_FETT = "Söhne Fett"                    # Heavy bold (rarely used)

# Path to the static brand assets directory — used for the cover logo.
_STATIC_IMG_DIR = Path(__file__).resolve().parent.parent / "static" / "img"


# ── Theme presets ───────────────────────────────────────────────────
# Two on-brand looks for the proposal: a white-paper LIGHT variant for
# print / classic advisory documents, and an on-screen DARK variant
# that mirrors the brand book's hero layouts (§7.1, Nightblue ground
# with Sandstone accents). The toggle applies to the whole document —
# cover, page chrome, headings, body, tables.

from dataclasses import dataclass as _dc


@_dc(frozen=True)
class Theme:
    """All theme-dependent colors and the logo PNG variant.

    Hex strings (without leading '#') for fields that feed
    ``_shade_cell`` (table fills, page background); RGBColor objects
    for font colors. Keep both forms so we don't double-convert.
    """
    name: str

    # Page-level
    page_bg_hex: str        # body page background ("FFFFFF" or "010626")
    cover_bg_hex: str       # cover background
    body_text: RGBColor     # default paragraph color
    muted_text: RGBColor    # subline / footnote color
    subline_text: RGBColor  # bigger subline (Söhne Leicht under titles)
    rule_color: RGBColor    # action-rule underline color

    # Headings + brand strip
    heading_color: RGBColor       # action-title color
    kicker_color: RGBColor        # small kicker above each section
    claim_color: RGBColor         # "The Digital Asset Boutique." on cover
    headline_color: RGBColor      # Big Headline on cover
    cover_subline_color: RGBColor # subline under cover headline
    running_strip_color: RGBColor # header/footer text strip

    # Tables
    table_header_fill_hex: str    # header row fill
    table_header_text: RGBColor   # header row text
    table_zebra_hex: str          # alternating-row fill
    table_body_text: RGBColor     # default cell text

    # Logo variant for cover (white logo on dark, dark logo on light).
    cover_logo_variant: str        # "white" or "dark"
    header_logo_variant: str       # logo in running header


LIGHT_THEME = Theme(
    name="light",
    page_bg_hex="FFFFFF",
    cover_bg_hex="FFFFFF",                         # white cover (brand §2.4 — dark logo on white)
    body_text=DEEP_INDIGO,
    muted_text=TEXT_MUTED,
    subline_text=TEXT_MUTED,
    rule_color=NIGHTBLUE,
    heading_color=NIGHTBLUE,
    kicker_color=SUNSET_EMBER,
    claim_color=NIGHTBLUE,                         # dark claim on white cover
    headline_color=NIGHTBLUE,
    cover_subline_color=DEEP_INDIGO,
    running_strip_color=TEXT_MUTED,
    table_header_fill_hex="010626",                # Nightblue header row
    table_header_text=WHITE,
    table_zebra_hex=SANDSTONE_50_HEX,
    table_body_text=DEEP_INDIGO,
    cover_logo_variant="dark",                     # dark logo on white cover
    header_logo_variant="dark",                    # dark logo on white body pages
)


DARK_THEME = Theme(
    name="dark",
    page_bg_hex="010626",                          # Nightblue body pages
    cover_bg_hex="010626",
    body_text=RGBColor(0xEC, 0xE8, 0xE5),          # Cream
    muted_text=SANDSTONE,
    subline_text=SANDSTONE,
    rule_color=SANDSTONE,
    heading_color=SANDSTONE,
    kicker_color=SUNSET_EMBER,                     # Sunset Ember pops on Nightblue
    claim_color=SANDSTONE,
    headline_color=SANDSTONE,
    cover_subline_color=WHITE,
    running_strip_color=SANDSTONE,
    table_header_fill_hex="0B688C",                # Electric Sky (sits on Nightblue body)
    table_header_text=WHITE,
    table_zebra_hex="060D43",                      # Deep Indigo
    table_body_text=RGBColor(0xEC, 0xE8, 0xE5),    # Cream
    cover_logo_variant="white",                    # white logo on Nightblue
    header_logo_variant="white",                   # white logo on Nightblue body
)


def _resolve_theme(ctx: dict) -> Theme:
    """Pick the proposal theme from the context dict. Accepts
    'light' or 'dark' (case-insensitive). Falls back to LIGHT when
    the value is missing or unknown — light is the brand-canonical
    print look and matches Jannick's WIP template.
    """
    name = str(ctx.get("theme") or "light").strip().lower()
    return DARK_THEME if name == "dark" else LIGHT_THEME


# ── Helpers ─────────────────────────────────────────────────────────


def _svg_to_png(svg: str, width_px: int = 720) -> Optional[bytes]:
    """Rasterise SVG to PNG bytes via cairosvg. Return None on failure
    so the renderer can drop the image silently rather than crash."""
    if not svg:
        return None
    try:
        import cairosvg  # type: ignore
        return cairosvg.svg2png(bytestring=svg.encode("utf-8"), output_width=width_px)
    except Exception:
        return None


def _apply_font_family(run_or_style_element, family_name: str) -> None:
    """Set ascii/hAnsi/cs font names on a run or style element via raw
    OOXML so Word respects the latin face for mixed-script paragraphs.
    python-docx's font.name only writes ``w:ascii`` by default.
    """
    rpr = run_or_style_element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for axis in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{axis}"), family_name)


def _set_default_font(doc: Document, theme: Theme = LIGHT_THEME) -> None:
    """Configure document defaults so Word picks brand fonts when
    available and Calibri / Cambria when not.

    Brand hierarchy from Short Brand Guideline §5.2:
        Bodytext  = Söhne Buch, 12pt / 16pt LH
        Subline   = Söhne Leicht, 22pt / 28pt LH
        Headline  = Sometimes Times Medium, 68pt / 78pt LH

    A4 advisory documents need denser layout than the brand poster
    examples, so we scale down: body 11pt, sublines 13pt, headings
    18-40pt. The font *family* names match the brand book exactly.
    """
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = theme.body_text
    _apply_font_family(style.element, FONT_BODY)

    # Heading styles map to the brand's Big Headline / Subline /
    # Special-heading tiers. We set the family name on each style so
    # Word picks the right weight without us forcing bold=True (which
    # would synthesise a fake bold on top of an already-weighted face).
    heading_specs = [
        (1, FONT_HEADING, 32, theme.heading_color),   # Cover / section dividers
        (2, FONT_HEADING, 18, theme.heading_color),   # Action titles
        (3, FONT_KRAFTIG, 12, theme.heading_color),   # Block-level subhead
    ]
    for level, family, size_pt, color in heading_specs:
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = family
        hs.font.color.rgb = color
        hs.font.size = Pt(size_pt)
        hs.font.bold = False
        _apply_font_family(hs.element, family)


def _logo_image_path(variant: str = "white") -> Optional[Path]:
    """Return the path to a pre-rendered Teroxx logo PNG.

    ``variant='white'`` returns the Sandstone/white logo for use on
    Nightblue backgrounds (cover). ``variant='dark'`` returns the
    Nightblue logo for use on light backgrounds (page header).

    We ship two rasterised PNGs in static/img/ so the renderer doesn't
    depend on cairosvg — which fails to load on some macOS setups
    (libcairo dynamic-lib mismatch). Returns None if the file is
    missing, so callers can skip embedding silently.
    """
    name = "logo-white.png" if variant == "white" else "logo-dark.png"
    path = _STATIC_IMG_DIR / name
    return path if path.exists() else None


def _page_setup(doc: Document, ctx: dict, theme: Theme = LIGHT_THEME) -> None:
    """A4 with brand-aligned margins, running header and running footer.

    Brand layout templates (Short Brand Guideline §7.1) put a small
    Söhne meta strip at the top of every spread:

        Teroxx — The Digital Asset Boutique. │ Section │ Date │ Page

    We mirror that as the Word "different first page" header (so the
    cover stays clean) and add a matching footer with just the page
    number on the right.
    """
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.0)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.1)
        # Brand cover should have no header chrome — Word respects
        # this via "different first page".
        section.different_first_page_header_footer = True

    lang = ctx.get("lang") or I18N_DEFAULT
    section_title = t("running.header_title", lang)
    client_name = (ctx.get("client") or {}).get("name") or ""
    prepared_date = ctx.get("prepared_date") or ""

    # Primary (non-first-page) header: brand strip on the left.
    sec = doc.sections[0]
    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Single tab stop on the right edge so we get "Brand · Section ...........  Date · Page"
    # page_width / margins are ints in EMU; convert to twips for w:pos.
    tab_pos_twips = (sec.page_width - sec.left_margin - sec.right_margin) // 635
    pPr = hp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(tab_pos_twips))  # twentieths-of-a-point
    tabs.append(tab)
    pPr.append(tabs)

    # Small Teroxx logo at the left of every body-page header, before
    # the brand text strip. Brand book §2.5 — minimum height 5mm
    # without protection zone; we use ~4mm here because the logo
    # repeats every page and the header strip must stay slim.
    header_logo = _logo_image_path(theme.header_logo_variant)
    if header_logo is not None:
        logo_run = hp.add_run()
        logo_run.add_picture(str(header_logo), height=Cm(0.4))
        hp.add_run("   ")  # small gap before the text strip

    left = hp.add_run("Teroxx — The Digital Asset Boutique.")
    left.font.size = Pt(8)
    left.font.color.rgb = theme.running_strip_color
    left.font.name = FONT_BODY
    _apply_font_family(left.element, FONT_BODY)

    sep1 = hp.add_run(f"   {section_title}")
    sep1.font.size = Pt(8)
    sep1.font.color.rgb = theme.running_strip_color
    sep1.font.name = FONT_LEICHT
    _apply_font_family(sep1.element, FONT_LEICHT)

    hp.add_run("\t")

    right = hp.add_run(f"{prepared_date}")
    right.font.size = Pt(8)
    right.font.color.rgb = theme.running_strip_color
    right.font.name = FONT_LEICHT
    _apply_font_family(right.element, FONT_LEICHT)

    # Page-number footer on the right, brand "X of Y" style.
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # client name on left, page X on right — single line, tabbed.
    pPr2 = fp._p.get_or_add_pPr()
    tabs2 = OxmlElement("w:tabs")
    tab2 = OxmlElement("w:tab")
    tab2.set(qn("w:val"), "right")
    tab2.set(qn("w:pos"), str(tab_pos_twips))
    tabs2.append(tab2)
    pPr2.append(tabs2)
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cname_run = fp.add_run(f"Confidential · {client_name}")
    cname_run.font.size = Pt(7.5)
    cname_run.font.color.rgb = theme.running_strip_color
    cname_run.font.name = FONT_LEICHT
    _apply_font_family(cname_run.element, FONT_LEICHT)

    fp.add_run("\t")
    # PAGE field code — Word renders the live page number.
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    page_run = fp.add_run()
    page_run.font.size = Pt(7.5)
    page_run.font.color.rgb = theme.running_strip_color
    page_run.font.name = FONT_LEICHT
    _apply_font_family(page_run.element, FONT_LEICHT)
    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)

    # First-page header/footer left empty — cover stays clean.
    first_h = sec.first_page_header.paragraphs[0]
    first_h.text = ""
    first_f = sec.first_page_footer.paragraphs[0]
    first_f.text = ""


def _shade_cell(cell, hex_color: str) -> None:
    """Apply a background fill to a table cell via raw OOXML.

    python-docx exposes no first-class cell-shading API, so we attach
    a ``<w:shd>`` element to the cell's tcPr. Hex without leading '#'.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: Optional[RGBColor] = None,
                   align: int = WD_ALIGN_PARAGRAPH.LEFT, size_pt: float = 10.0,
                   font: str = FONT_BODY) -> None:
    cell.text = ""  # clear default empty paragraph contents
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = font
    _apply_font_family(run.element, font)
    run.font.size = Pt(size_pt)
    # ``bold=True`` on a Söhne-Buch run forces Word to synthesise a fake
    # bold, which looks subtly wrong. Caller should instead pass
    # font=FONT_HALBFETT or font=FONT_KRAFTIG when emphasis is needed.
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_horizontal_rule(doc: Document, color: RGBColor = NIGHTBLUE) -> None:
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color}")  # RGBColor str = 'RRGGBB'
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_md_block(doc: Document, md: str) -> None:
    """Render a small subset of markdown into Word paragraphs.

    Supports bullet lists, numbered lists, paragraphs separated by
    blank lines, and inline **bold** / *italic*. Not a full markdown
    renderer — keeps the artifact editable in Word with simple
    structure. Anything fancier should be drafted directly in Word
    after generation.
    """
    if not md:
        return
    blocks: list[list[str]] = []
    current: list[str] = []
    for raw in md.replace("\r\n", "\n").split("\n"):
        if raw.strip() == "":
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(raw)
    if current:
        blocks.append(current)

    bullet_re = re.compile(r"^\s*[-*]\s+(.*)$")
    numbered_re = re.compile(r"^\s*\d+\.\s+(.*)$")

    for block in blocks:
        if all(bullet_re.match(line) for line in block):
            for line in block:
                doc.add_paragraph(bullet_re.match(line).group(1), style="List Bullet")
            continue
        if all(numbered_re.match(line) for line in block):
            for line in block:
                doc.add_paragraph(numbered_re.match(line).group(1), style="List Number")
            continue
        text = " ".join(line.strip() for line in block)
        p = doc.add_paragraph()
        _add_inline_formatted_runs(p, text)


_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _add_inline_formatted_runs(paragraph, text: str) -> None:
    for chunk in _INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = paragraph.add_run(chunk[2:-2])
            r.font.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            r = paragraph.add_run(chunk[1:-1])
            r.font.italic = True
        else:
            paragraph.add_run(chunk)


def _T(ctx: dict, key: str, **fmt) -> str:
    """Shorthand: pull the bound translator off the ctx dict."""
    lang = ctx.get("lang") or I18N_DEFAULT
    return t(key, lang, **fmt)


def _money(ctx: dict, amount: float) -> str:
    """Currency-aware formatter. Reads client.currency and lang from ctx.

    USD → "$50,000"
    EUR + lang=de → "50.000 €" (German thousand-separator convention)
    EUR + lang=en → "€50,000"
    Other → "{X,XXX} {ccy}"
    """
    if amount is None:
        return "-"
    ccy = (ctx.get("client", {}).get("currency") or "USD").upper()
    lang = ctx.get("lang") or I18N_DEFAULT
    val = float(amount)
    if ccy == "USD":
        return f"${val:,.0f}"
    if ccy == "EUR":
        if lang == "de":
            s = f"{val:,.0f}".replace(",", ".")
            return f"{s} €"
        return f"€{val:,.0f}"
    return f"{val:,.0f} {ccy}"


def _country_label(client: dict, lang: str) -> str:
    """Best-effort country display string.

    Prefer the long ``domicile`` ("Berlin, DE") but fall back to the
    ISO ``domicile_country`` if the long form is missing.
    """
    return (client.get("domicile") or client.get("domicile_country") or "").strip()


def _md_or_placeholder(doc: Document, ctx: dict, md_value: str, placeholder_key: str) -> None:
    """Drop ``md_value`` as a paragraph block, or a muted placeholder
    instructing the advisor to fill the section in. The placeholder is
    italicised so it's visually obvious it needs editing before sending."""
    if md_value:
        _add_md_block(doc, md_value)
        return
    p = doc.add_paragraph()
    r = p.add_run(_T(ctx, placeholder_key).strip("_ ").strip())
    r.font.italic = True
    r.font.color.rgb = TEXT_MUTED
    r.font.size = Pt(10)


# ── Section builders ────────────────────────────────────────────────


def _set_section_page_color(section, hex_color: str) -> None:
    """Tint the whole page background of a Word section.

    Word doesn't expose a per-section page-color API; we attach a
    ``<w:background>`` to the document and ``<w:displayBackgroundShape/>``
    so the cover renders Nightblue on screen and (depending on print
    settings) on paper.
    """
    doc_el = section.part.document.element
    body = doc_el
    # The background element must live on the <w:document> root.
    root = body.getparent() if body.tag.endswith("body") else body
    bg = root.find(qn("w:background"))
    if bg is None:
        bg = OxmlElement("w:background")
        bg.set(qn("w:color"), hex_color)
        root.insert(0, bg)


def _cover(doc: Document, ctx: dict, theme: Theme = LIGHT_THEME) -> None:
    """Brand-aligned cover (Short Brand Guideline §7.1).

    Layout:
        Top-left:    white Teroxx logo (the primary brand element)
        Body:        "Allocation Proposal" — Sometimes Times Medium ~40pt
        Subline:     client name + profile — Söhne Leicht in Sandstone
        Brand claim: "The Digital Asset Boutique." — Sometimes Times
                     Medium in Sandstone (per §2.4 / §7.1)
        Footer:      confidentiality strip in muted Sandstone

    Detailed client / advisor metadata lives in the dedicated Client
    Information section that follows the welcome page — matches
    Jannick's template flow.
    """
    sec = doc.sections[0]
    # Set the document-wide page color. Word's <w:background> is
    # document-wide (not section-scoped), so this also colors body
    # pages — which is exactly what we want for the dark theme. For
    # the light theme the cover still needs Nightblue but body pages
    # need white; we work around that by always painting the cover
    # background via the page color *and* relying on Word's default
    # white page-fill for body pages (set via "displayBackgroundShape"
    # being absent for screen viewers; for print the cover image
    # carries the dark fill itself).
    _set_section_page_color(sec, theme.page_bg_hex)

    # Pull cover layout inward so the logo + headline breathe.
    cover_para = doc.add_paragraph()
    cover_para.paragraph_format.space_before = Pt(6)

    # ── Logo (top-left, variant chosen by theme) ──
    logo_path = _logo_image_path(theme.cover_logo_variant)
    if logo_path is not None:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.add_run().add_picture(str(logo_path), width=Cm(4.2))

    # Spacer to push the headline into the lower-middle of the page.
    for _ in range(6):
        doc.add_paragraph()

    # ── Big Headline (Sometimes Times Medium, theme color) ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(_T(ctx, "cover.title"))
    title_run.font.size = Pt(40)
    title_run.font.bold = False
    title_run.font.color.rgb = theme.headline_color
    title_run.font.name = FONT_HEADING
    _apply_font_family(title_run.element, FONT_HEADING)

    # ── Subline (Söhne Leicht, theme color) ──
    sub_p = doc.add_paragraph()
    sub_text = ctx.get("cover_subtitle") or ""
    if sub_text:
        sub_run = sub_p.add_run(sub_text)
        sub_run.font.size = Pt(13)
        sub_run.font.color.rgb = theme.cover_subline_color
        sub_run.font.name = FONT_LEICHT
        _apply_font_family(sub_run.element, FONT_LEICHT)

    for _ in range(5):
        doc.add_paragraph()

    # ── Brand claim "The Digital Asset Boutique." ──
    # Section 7.1 of the brand book: "The Digital Asset Boutique is
    # always set in Sometimes Times Medium and the color is selected
    # in relation to the background: on Nightblue the claim is set
    # in Sandstone."
    claim_p = doc.add_paragraph()
    claim_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    claim_run = claim_p.add_run(_T(ctx, "cover.brand_claim"))
    claim_run.font.size = Pt(22)
    claim_run.font.color.rgb = theme.claim_color
    claim_run.font.name = FONT_HEADING
    _apply_font_family(claim_run.element, FONT_HEADING)

    for _ in range(2):
        doc.add_paragraph()

    # ── Confidentiality strip (very small Söhne Leicht) ──
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.LEFT
    foot_run = foot.add_run(_T(ctx, "cover.confidential"))
    foot_run.font.size = Pt(8.5)
    foot_run.font.color.rgb = theme.muted_text
    foot_run.font.name = FONT_LEICHT
    _apply_font_family(foot_run.element, FONT_LEICHT)

    doc.add_page_break()


def _welcome(doc: Document, ctx: dict) -> None:
    """Personal salutation + welcome paragraphs.

    Directional from Jannick's template: every IA proposal opens with
    "Sehr geehrter Herr {Lastname}," and a four-paragraph welcome.
    Advisor can fully override via the salutation / welcome_md fields
    on the client record; defaults are provided so the auto-generated
    DOCX still looks finished out of the box.
    """
    lang = ctx.get("lang", I18N_DEFAULT)
    ov = ctx.get("overrides") or {}
    client = ctx.get("client", {})

    # Salutation: explicit override > default with client name.
    salutation = ov.get("salutation") or _T(
        ctx, "welcome.salutation_default",
        name=client.get("name", "")
    )
    p = doc.add_paragraph()
    pr = p.add_run(salutation)
    pr.font.size = Pt(11)

    welcome_body = ov.get("welcome_md") or _T(ctx, "welcome.body_default")
    _add_md_block(doc, welcome_body)
    doc.add_page_break()


def _client_info(doc: Document, ctx: dict) -> None:
    """Two stacked tables: client information + risk profile.

    Mirrors Jannick's tables 1 + 2. Status level and consultation
    date come from per-client overrides; risk-tolerance defaults are
    derived from the configured profile when no override is set.
    """
    lang = ctx.get("lang", I18N_DEFAULT)
    client = ctx.get("client", {})
    ov = ctx.get("overrides") or {}

    _section_header(doc, ctx, "client_info.heading", _T(ctx, "client_info.heading"))

    # ── Client information table ──
    rows = [
        (_T(ctx, "client_info.prepared_by"), _T(ctx, "client_info.prepared_by_team")),
        (
            _T(ctx, "client_info.consultation_date"),
            ov.get("consultation_date") or ctx.get("prepared_date", ""),
        ),
        (_T(ctx, "client_info.client_name"), client.get("name", "")),
        (_T(ctx, "client_info.country"), _country_label(client, lang)),
    ]
    if ov.get("status_level"):
        rows.append((_T(ctx, "client_info.status_level"), ov["status_level"]))

    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(6.0)
    table.columns[1].width = Cm(10.5)
    for i, (lbl, val) in enumerate(rows):
        _set_cell_text(table.rows[i].cells[0], lbl, bold=True, color=TEXT_MUTED, size_pt=9)
        _set_cell_text(table.rows[i].cells[1], val, size_pt=10.5)
        if i % 2 == 0:
            _shade_cell(table.rows[i].cells[0], "F6F4F0")
            _shade_cell(table.rows[i].cells[1], "F6F4F0")

    doc.add_paragraph()

    # ── Risk profile mini-table ──
    h = doc.add_paragraph()
    hr = h.add_run(_T(ctx, "risk_profile.heading"))
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = NIGHTBLUE

    tolerance = t(
        f"risk_profile.tolerance_default.{client.get('profile', '')}",
        lang,
    ) or profile_label(client.get("profile", ""), lang)
    risk_rows = [
        (_T(ctx, "risk_profile.tolerance"), tolerance),
        (_T(ctx, "risk_profile.horizon"), _T(ctx, "risk_profile.horizon_default")),
        (_T(ctx, "risk_profile.objective"), _T(ctx, "risk_profile.objective_default")),
    ]
    rtable = doc.add_table(rows=len(risk_rows), cols=2)
    rtable.autofit = False
    rtable.columns[0].width = Cm(6.0)
    rtable.columns[1].width = Cm(10.5)
    for i, (lbl, val) in enumerate(risk_rows):
        _set_cell_text(rtable.rows[i].cells[0], lbl, bold=True, color=TEXT_MUTED, size_pt=9)
        _set_cell_text(rtable.rows[i].cells[1], val, size_pt=10.5)
        if i % 2 == 0:
            _shade_cell(rtable.rows[i].cells[0], "F6F4F0")
            _shade_cell(rtable.rows[i].cells[1], "F6F4F0")

    doc.add_page_break()


def _consultation_summary(doc: Document, ctx: dict) -> None:
    """Advisor's account of the meeting + agreed positioning.

    Maps to Jannick's "Zusammenfassung der Beratung / des Gesprächs"
    heading. Always rendered so the advisor sees the slot to fill in;
    when no summary_md is supplied a muted-italic placeholder appears
    instructing them to describe what was discussed in the call
    (Leonardo, 2026-05-19).
    """
    ov = ctx.get("overrides") or {}
    body = (ov.get("summary_md") or "").strip()

    _section_header(doc, ctx, "page.consultation", _T(ctx, "page.consultation"))
    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "overrides.summary_sub"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED
    sr.font.name = FONT_LEICHT
    _apply_font_family(sr.element, FONT_LEICHT)

    _md_or_placeholder(doc, ctx, body, "consultation.placeholder")
    doc.add_page_break()


def _market_analysis(doc: Document, ctx: dict) -> None:
    """Long-form market commentary slot — the section Jannick invests
    the most time in. Renders the advisor's market_analysis_md, or a
    visible italic placeholder if empty (advisor knows what to fill).
    """
    _section_header(doc, ctx, "page.market_analysis", _T(ctx, "page.market_analysis"))
    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "market_analysis.subheading"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED
    ov = ctx.get("overrides") or {}
    _md_or_placeholder(doc, ctx, ov.get("market_analysis_md", ""), "market_analysis.placeholder")
    doc.add_page_break()


def _portfolio_detail(doc: Document, ctx: dict) -> None:
    """Combined portfolio section: parameters → asset-class summary →
    per-ticker target weights → tier-bar exhibit → donut exhibit.

    Folds the prior _exec_summary KPI strip into the parameters table
    here so the proposal reads as a single coherent "your portfolio"
    block rather than two thinly-separated pages."""
    _section_header(doc, ctx, "page.portfolio_detail", ctx.get("allocation_title", _T(ctx, "page.your_new_portfolio")))

    client = ctx.get("client", {})
    lang = ctx.get("lang", I18N_DEFAULT)

    # ── Parameters block (Jannick table 3) ──
    # Thematic baskets carry no defensive sleeve, so that row is replaced
    # by the basket's weighting basis (market-cap vs fundamental score).
    if ctx.get("is_basket"):
        wl_key = ("weighting.market_cap"
                  if ctx.get("weighting_label") == "market-cap"
                  else "weighting.fundamental")
        sleeve_or_weighting = (_T(ctx, "kpi.weighting"), _T(ctx, wl_key))
    else:
        sleeve_or_weighting = (_T(ctx, "kpi.defensive_sleeve"), f"{ctx.get('defensive_pct', 0):.0f}%")
    params = [
        (_T(ctx, "kpi.risk_profile"), profile_label(client.get("profile", ""), lang)),
        (_T(ctx, "table.asset"), ctx.get("universe", "")),
        (_T(ctx, "kpi.portfolio_value"), _money(ctx, ctx.get("portfolio_value", 0))),
        sleeve_or_weighting,
        (_T(ctx, "kpi.positions"), str(ctx.get("allocation_count", 0))),
    ]
    h = doc.add_paragraph()
    hr = h.add_run(_T(ctx, "exhibit.what_this_means"))
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = NIGHTBLUE
    ptable = doc.add_table(rows=len(params), cols=2)
    ptable.autofit = False
    ptable.columns[0].width = Cm(6.0)
    ptable.columns[1].width = Cm(10.5)
    for i, (lbl, val) in enumerate(params):
        _set_cell_text(ptable.rows[i].cells[0], lbl, bold=True, color=TEXT_MUTED, size_pt=9)
        _set_cell_text(ptable.rows[i].cells[1], val, size_pt=10.5)
        if i % 2 == 0:
            _shade_cell(ptable.rows[i].cells[0], "F6F4F0")
            _shade_cell(ptable.rows[i].cells[1], "F6F4F0")

    doc.add_paragraph()

    # ── "What this means" bullets ──
    bullets = ctx.get("exec_bullets") or []
    if bullets:
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")
        doc.add_paragraph()

    # ── Asset-class aggregation table (directional from Jannick) ──
    ac_rows = ctx.get("asset_class_rows") or []
    nonzero = [r for r in ac_rows if (r.get("share_pct") or 0) > 0]
    if nonzero:
        h2 = doc.add_paragraph()
        hr2 = h2.add_run(_T(ctx, "assetclass.heading"))
        hr2.font.bold = True
        hr2.font.size = Pt(12)
        hr2.font.color.rgb = NIGHTBLUE
        sub2 = doc.add_paragraph()
        sr2 = sub2.add_run(_T(ctx, "assetclass.subheading"))
        sr2.font.size = Pt(9)
        sr2.font.italic = True
        sr2.font.color.rgb = TEXT_MUTED
        ac_table = doc.add_table(rows=1 + len(nonzero) + 1, cols=2)
        headers = [_T(ctx, "assetclass.col_class"), _T(ctx, "assetclass.col_share")]
        for i, hd in enumerate(headers):
            _set_cell_text(ac_table.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
            _shade_cell(ac_table.rows[0].cells[i], "010626")
        total = 0.0
        for i, row in enumerate(nonzero, start=1):
            cells = ac_table.rows[i].cells
            _set_cell_text(cells[0], row["label"], size_pt=10)
            _set_cell_text(cells[1], f"{row['share_pct']:.1f}%",
                           align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=10)
            total += float(row.get("share_pct") or 0)
            if i % 2 == 0:
                for cell in cells:
                    _shade_cell(cell, "F6F4F0")
        tcells = ac_table.rows[-1].cells
        _set_cell_text(tcells[0], _T(ctx, "assetclass.total"), bold=True, size_pt=10)
        _set_cell_text(tcells[1], f"{total:.1f}%", bold=True,
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=10)
        _shade_cell(tcells[0], "BFB3A8")
        _shade_cell(tcells[1], "BFB3A8")

        doc.add_paragraph()

    # ── Donut exhibit ──
    png = _svg_to_png(ctx.get("donut_svg") or "", width_px=600)
    if png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(png), width=Cm(10))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(_T(ctx, "exhibit.recommended_alloc"))
        cr.font.size = Pt(8.5)
        cr.font.italic = True
        cr.font.color.rgb = TEXT_MUTED

    # No page break: _allocation_table continues this same merged
    # "your portfolio in detail" section straight below, so the
    # allocation finding sentence is not repeated on a second header.
    doc.add_paragraph()


def _wishes_section(doc: Document, ctx: dict) -> None:
    """Optional client-wishes section. Rendered only when the advisor
    has explicitly drafted it; no placeholder so the auto-PDF stays
    quiet for clients without specific requests."""
    ov = ctx.get("overrides") or {}
    body = ov.get("wishes_md", "")
    if not (body and body.strip()):
        return
    _section_header(doc, ctx, "page.wishes", _T(ctx, "page.wishes"))
    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "overrides.wishes_sub"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED
    _add_md_block(doc, body)
    doc.add_page_break()


def _dca_table(doc: Document, ctx: dict) -> None:
    """Monthly Buy Schedule — the phased-build DCA table.

    Rendered as a sub-block of the merged "your portfolio in detail"
    section, directly under the recommended allocation, and only when
    the client has a recurring-buy plan (``dca_rows`` present)."""
    rows = ctx.get("dca_rows") or []
    if not rows:
        return
    meta = ctx.get("dca_meta") or {}
    doc.add_paragraph()
    h = doc.add_paragraph()
    hr = h.add_run(_T(ctx, "exhibit.phased_build"))
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = NIGHTBLUE
    sub2 = doc.add_paragraph()
    sr2 = sub2.add_run(_T(ctx, "exhibit.phased_build_sub",
                          horizon_months=meta.get("horizon_months", 0)))
    sr2.font.size = Pt(9.5)
    sr2.font.italic = True
    sr2.font.color.rgb = TEXT_MUTED

    table = doc.add_table(rows=1 + len(rows), cols=4)
    headers = [
        _T(ctx, "table.asset"),
        _T(ctx, "table.weight_pct"),
        _T(ctx, "table.monthly_buy"),
        _T(ctx, "table.horizon_total"),
    ]
    for i, hd in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
        _shade_cell(table.rows[0].cells[i], "010626")
    for i, r in enumerate(rows, start=1):
        cells = table.rows[i].cells
        c = cells[0]
        c.text = ""
        p = c.paragraphs[0]
        rt = p.add_run(r.get("ticker", ""))
        rt.font.bold = True
        rt.font.size = Pt(10)
        rn = p.add_run(f"  {r.get('name', '')}")
        rn.font.size = Pt(9)
        rn.font.color.rgb = TEXT_MUTED
        _set_cell_text(cells[1], f"{(r.get('portfolio_pct', 0) * 100):.2f}%",
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[2], _money(ctx, r.get("monthly_buy", 0)),
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[3], _money(ctx, r.get("horizon_total", 0)),
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        if i % 2 == 0:
            for cell in cells:
                _shade_cell(cell, "F6F4F0")


def _strategy_section(doc: Document, ctx: dict) -> None:
    """Strategy section — the advisor's drafted execution-plan narrative.

    The DCA Monthly Buy Schedule table now renders inside the merged
    portfolio section (see _dca_table); this section carries only the
    drafted execution plan. Falls back to PAM's default implementation
    paragraph when no plan is drafted, so the page is never blank."""
    ov = ctx.get("overrides") or {}
    plan_md = ov.get("execution_plan_md", "")
    if not plan_md:
        # Auto-generated implementation paragraph + review cadence.
        _implementation_section(doc, ctx)
        return

    _section_header(doc, ctx, "page.strategy", _T(ctx, "exhibit.phased_build"))
    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "overrides.execution_sub"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED
    _add_md_block(doc, plan_md)
    doc.add_page_break()


def _fazit_section(doc: Document, ctx: dict) -> None:
    """Conclusion / Fazit. Optional; placeholder absent so empty
    proposals don't carry a stub heading."""
    ov = ctx.get("overrides") or {}
    body = ov.get("conclusion_md", "")
    if not (body and body.strip()):
        return
    _section_header(doc, ctx, "page.fazit", _T(ctx, "page.fazit"))
    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "fazit.subheading"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED
    _add_md_block(doc, body)
    doc.add_page_break()


def _fees_section(doc: Document, ctx: dict) -> None:
    """Fee structure table (directional from Jannick template table 6).

    Renders only when the advisor supplied at least one fee row; the
    table headers/cells are translated, but the values come verbatim
    from the override so the advisor can phrase exemptions, %s, or
    explanatory notes however the engagement requires.
    """
    ov = ctx.get("overrides") or {}
    fees = ov.get("fees") or []
    if not isinstance(fees, list) or not fees:
        return
    cleaned = [
        f for f in fees
        if isinstance(f, dict) and (f.get("name") or f.get("value"))
    ]
    if not cleaned:
        return

    _section_header(doc, ctx, "page.fees", _T(ctx, "page.fees"))
    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "fees.subheading"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    table = doc.add_table(rows=1 + len(cleaned), cols=2)
    headers = [_T(ctx, "fees.col_component"), _T(ctx, "fees.col_value")]
    for i, hd in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
        _shade_cell(table.rows[0].cells[i], "010626")
    for i, f in enumerate(cleaned, start=1):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], f.get("name", ""), size_pt=10)
        _set_cell_text(cells[1], f.get("value", ""), size_pt=10)
        if i % 2 == 0:
            for cell in cells:
                _shade_cell(cell, "F6F4F0")

    doc.add_page_break()


def _current_allocation(doc: Document, ctx: dict) -> None:
    """Current crypto allocation shown to the client BEFORE the new
    "Your portfolio in detail" section (Leonardo, 2026-05-19).

    Behaviour:
      * If ``ctx['current_allocation']`` is supplied (list of dicts with
        ticker / name / weight_pct / value), render it as a table.
      * Else if ``ctx['review']['pnl']['ticker_summary']`` has rows
        (existing-client flow), fall back to those.
      * Else render an editable 4-row placeholder table so the advisor
        can paste / type the client's current holdings directly in Word.

    The section is always rendered: even when the client has no crypto
    on file, advisors need the slot visible so they don't forget to
    fill it in before sending.
    """
    rows = ctx.get("current_allocation") or []
    if not rows:
        review = ctx.get("review") or {}
        pnl = review.get("pnl") or {}
        rows = pnl.get("ticker_summary") or pnl.get("by_ticker") or []

    _section_header(doc, ctx, "page.current_allocation",
                    _T(ctx, "page.current_allocation"))

    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "current_allocation.subheading"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED
    sr.font.name = FONT_LEICHT
    _apply_font_family(sr.element, FONT_LEICHT)

    if rows:
        # Real data: 3 columns. Keep it simple — clients see "what I own
        # today", not full P&L (that's in the review-flow section).
        table = doc.add_table(rows=1 + len(rows), cols=3)
        headers = [
            _T(ctx, "table.asset"),
            _T(ctx, "table.weight_now"),
            _T(ctx, "table.market_value"),
        ]
        for i, hd in enumerate(headers):
            _set_cell_text(table.rows[0].cells[i], hd, bold=True,
                           color=WHITE, size_pt=9.5, font=FONT_HALBFETT)
            _shade_cell(table.rows[0].cells[i], "010626")
        for i, r in enumerate(rows, start=1):
            cells = table.rows[i].cells
            c = cells[0]
            c.text = ""
            p = c.paragraphs[0]
            rt = p.add_run(r.get("ticker", ""))
            rt.font.bold = True
            rt.font.size = Pt(10)
            rt.font.name = FONT_HALBFETT
            _apply_font_family(rt.element, FONT_HALBFETT)
            rn = p.add_run(f"  {r.get('name', '')}")
            rn.font.size = Pt(9)
            rn.font.color.rgb = TEXT_MUTED
            _set_cell_text(
                cells[1],
                f"{r.get('weight_pct', r.get('weight', 0)):.1f}%",
                align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5,
            )
            _set_cell_text(
                cells[2],
                _money(ctx, r.get("value", r.get("current_value", 0))),
                align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5,
            )
            if i % 2 == 0:
                for cell in cells:
                    _shade_cell(cell, SANDSTONE_50_HEX)
    else:
        # Empty placeholder — 4 blank rows the advisor fills in directly.
        ph = doc.add_paragraph()
        pr = ph.add_run(_T(ctx, "current_allocation.placeholder"))
        pr.font.italic = True
        pr.font.color.rgb = TEXT_MUTED
        pr.font.size = Pt(10)
        pr.font.name = FONT_LEICHT
        _apply_font_family(pr.element, FONT_LEICHT)

        table = doc.add_table(rows=5, cols=3)
        headers = [
            _T(ctx, "table.asset"),
            _T(ctx, "table.weight_now"),
            _T(ctx, "table.market_value"),
        ]
        for i, hd in enumerate(headers):
            _set_cell_text(table.rows[0].cells[i], hd, bold=True,
                           color=WHITE, size_pt=9.5, font=FONT_HALBFETT)
            _shade_cell(table.rows[0].cells[i], "010626")
        for i in range(1, 5):
            cells = table.rows[i].cells
            _set_cell_text(cells[0], " ", size_pt=10)
            _set_cell_text(cells[1], " ", align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=10)
            _set_cell_text(cells[2], " ", align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=10)
            if i % 2 == 0:
                for cell in cells:
                    _shade_cell(cell, SANDSTONE_50_HEX)

    doc.add_page_break()


# ── Review-flow sections (existing-client proposals) ───────────────


def _current_holdings(doc: Document, ctx: dict) -> None:
    """Current portfolio: KPI strip (cost basis / market value / P&L /
    days held) + per-ticker holdings table with live mark-to-market."""
    review = ctx.get("review") or {}
    pnl = review.get("pnl") or {}
    rows = pnl.get("ticker_summary") or pnl.get("by_ticker") or []
    if not rows:
        # Fall back to lot-level rows if the ticker rollup is missing.
        rows = pnl.get("rows") or []
    summary = pnl.get("summary") or {}
    if not rows and not summary:
        # Nothing to render — likely client has no lots; the renderer
        # will switch the audience back to the new-client flow upstream
        # but we still guard here.
        return

    _section_header(
        doc, ctx, "page.current_portfolio",
        _T(ctx, "current_portfolio.title_with_total",
           value=_money(ctx, summary.get("total_value", 0))),
    )
    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "current_portfolio.subheading"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    # KPI strip — uses the same shape as portfolio_detail's params block.
    kpis = [
        (_T(ctx, "kpi.total_value"), _money(ctx, summary.get("total_value", 0))),
        (_T(ctx, "kpi.total_cost"), _money(ctx, summary.get("total_cost", 0))),
        (_T(ctx, "kpi.total_pnl"),
         f"{_money(ctx, summary.get('total_pnl', 0))} ({summary.get('total_pnl_pct', 0):.1f}%)"),
        (_T(ctx, "kpi.days_held"), str(summary.get("days_since_entry", 0))),
    ]
    ktable = doc.add_table(rows=len(kpis), cols=2)
    ktable.autofit = False
    ktable.columns[0].width = Cm(7.0)
    ktable.columns[1].width = Cm(9.5)
    for i, (lbl, val) in enumerate(kpis):
        _set_cell_text(ktable.rows[i].cells[0], lbl, bold=True, color=TEXT_MUTED, size_pt=9)
        _set_cell_text(ktable.rows[i].cells[1], val, size_pt=10.5)
        if i % 2 == 0:
            _shade_cell(ktable.rows[i].cells[0], "F6F4F0")
            _shade_cell(ktable.rows[i].cells[1], "F6F4F0")

    doc.add_paragraph()

    # Per-ticker holdings table.
    headers = [
        _T(ctx, "table.asset"),
        _T(ctx, "table.qty"),
        _T(ctx, "table.entry_price"),
        _T(ctx, "table.live_price"),
        _T(ctx, "table.market_value"),
        _T(ctx, "table.pnl_unrealized"),
        _T(ctx, "table.weight_now"),
    ]
    htable = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, hd in enumerate(headers):
        _set_cell_text(htable.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
        _shade_cell(htable.rows[0].cells[i], "010626")
    for i, r in enumerate(rows, start=1):
        cells = htable.rows[i].cells
        c = cells[0]
        c.text = ""
        p = c.paragraphs[0]
        rt = p.add_run(r.get("ticker", ""))
        rt.font.bold = True
        rt.font.size = Pt(10)
        rn = p.add_run(f"  {r.get('name', '')}")
        rn.font.size = Pt(9)
        rn.font.color.rgb = TEXT_MUTED
        _set_cell_text(cells[1], f"{r.get('quantity', 0):,.4f}".rstrip("0").rstrip("."),
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[2], _money(ctx, r.get("avg_entry_price", r.get("entry_price", 0))),
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[3], _money(ctx, r.get("current_price", 0)),
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[4], _money(ctx, r.get("current_value", 0)),
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        pnl_val = r.get("pnl", 0)
        pnl_text = f"{_money(ctx, pnl_val)} ({r.get('pnl_pct', 0):.1f}%)"
        # Colour the P&L cell green / red so the advisor can scan the
        # row at a glance even before reading the value.
        pnl_cell = cells[5]
        pnl_cell.text = ""
        pp = pnl_cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pr = pp.add_run(pnl_text)
        pr.font.size = Pt(9.5)
        if pnl_val > 0:
            pr.font.color.rgb = GAIN
        elif pnl_val < 0:
            pr.font.color.rgb = LOSS
        _set_cell_text(cells[6], f"{r.get('weight', 0):.1f}%",
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        if i % 2 == 0:
            for cell in cells:
                _shade_cell(cell, "F6F4F0")

    doc.add_page_break()


def _drift_analysis(doc: Document, ctx: dict) -> None:
    """Per-position drift vs the configured profile target."""
    review = ctx.get("review") or {}
    drift = review.get("drift") or {}
    rows = drift.get("rows") or []
    if not rows:
        return

    lang = ctx.get("lang", I18N_DEFAULT)
    profile = ctx.get("client", {}).get("profile", "")
    threshold = drift.get("threshold_pp", 0)

    _section_header(doc, ctx, "page.drift_analysis", _T(ctx, "page.drift_analysis"))

    sub = doc.add_paragraph()
    sr = sub.add_run(_T(
        ctx, "drift_analysis.subheading",
        profile=profile_label(profile, lang),
        threshold=f"{threshold:.0f}",
    ))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    # Callout for the largest drift.
    max_ticker = drift.get("max_drift_ticker")
    max_pp = drift.get("max_drift_pp", 0)
    if max_ticker:
        status_key = "drift_analysis.attention_yes" if drift.get("attention") else "drift_analysis.attention_no"
        callout = doc.add_paragraph()
        cr = callout.add_run(_T(
            ctx, "drift_analysis.max_drift_callout",
            ticker=max_ticker, drift=f"{max_pp:.1f}",
            status=_T(ctx, status_key),
        ))
        cr.font.size = Pt(10)
        cr.font.bold = True
        cr.font.color.rgb = SUNSET_EMBER if drift.get("attention") else TEXT_BODY

    headers = [
        _T(ctx, "table.asset"),
        _T(ctx, "table.weight_now"),
        _T(ctx, "table.weight_target"),
        _T(ctx, "table.drift_pp"),
    ]
    dtable = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, hd in enumerate(headers):
        _set_cell_text(dtable.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
        _shade_cell(dtable.rows[0].cells[i], "010626")
    for i, r in enumerate(rows, start=1):
        cells = dtable.rows[i].cells
        _set_cell_text(cells[0], r.get("ticker", ""), bold=True, size_pt=10)
        _set_cell_text(cells[1], f"{r.get('current_pct', 0):.2f}%",
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[2], f"{r.get('target_pct', 0):.2f}%",
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        drift_pp = float(r.get("drift_pp", 0) or 0)
        sign = "+" if drift_pp > 0 else ""
        drift_cell = cells[3]
        drift_cell.text = ""
        dp = drift_cell.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dr = dp.add_run(f"{sign}{drift_pp:.2f}pp")
        dr.font.size = Pt(9.5)
        dr.font.bold = abs(drift_pp) >= float(threshold or 0)
        if abs(drift_pp) >= float(threshold or 0):
            dr.font.color.rgb = SUNSET_EMBER
        if i % 2 == 0:
            for cell in cells:
                _shade_cell(cell, "F6F4F0")

    doc.add_page_break()


def _rebalance_actions(doc: Document, ctx: dict) -> None:
    """Recommended BUY/SELL trades to bring the portfolio to target."""
    review = ctx.get("review") or {}
    rebal = review.get("rebalance") or {}
    rows = rebal.get("rows") or []
    actionable = [r for r in rows if (r.get("action") or "") and abs(float(r.get("difference") or 0)) > 1]
    lang = ctx.get("lang", I18N_DEFAULT)
    profile = ctx.get("client", {}).get("profile", "")
    net = rebal.get("net_rebalance", 0)

    _section_header(doc, ctx, "page.rebalance", _T(ctx, "page.rebalance"))

    sub = doc.add_paragraph()
    sr = sub.add_run(_T(
        ctx, "rebalance.subheading",
        profile=profile_label(profile, lang),
        net=_money(ctx, net),
    ))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    if not actionable:
        p = doc.add_paragraph()
        pr = p.add_run(_T(ctx, "rebalance.no_action"))
        pr.font.italic = True
        pr.font.color.rgb = TEXT_MUTED
        pr.font.size = Pt(10.5)
        doc.add_page_break()
        return

    headers = [
        _T(ctx, "table.asset"),
        _T(ctx, "table.weight_target"),
        _T(ctx, "table.market_value"),
        _T(ctx, "table.delta_usd"),
        _T(ctx, "table.action"),
    ]
    rtable = doc.add_table(rows=1 + len(actionable), cols=len(headers))
    for i, hd in enumerate(headers):
        _set_cell_text(rtable.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
        _shade_cell(rtable.rows[0].cells[i], "010626")
    for i, r in enumerate(actionable, start=1):
        cells = rtable.rows[i].cells
        c = cells[0]
        c.text = ""
        p = c.paragraphs[0]
        rt = p.add_run(r.get("ticker", ""))
        rt.font.bold = True
        rt.font.size = Pt(10)
        rn = p.add_run(f"  {r.get('name', '')}")
        rn.font.size = Pt(9)
        rn.font.color.rgb = TEXT_MUTED
        _set_cell_text(cells[1], f"{(r.get('target_pct') or 0) * 100:.2f}%",
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[2], _money(ctx, r.get("current_usd", 0)),
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        delta = float(r.get("difference") or 0)
        sign = "+" if delta > 0 else ""
        _set_cell_text(cells[3], f"{sign}{_money(ctx, delta)}",
                       align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        action_cell = cells[4]
        action_cell.text = ""
        ap = action_cell.paragraphs[0]
        ap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ar = ap.add_run(r.get("action", ""))
        ar.font.bold = True
        ar.font.size = Pt(9.5)
        if r.get("action") == "BUY":
            ar.font.color.rgb = GAIN
        elif r.get("action") == "SELL":
            ar.font.color.rgb = LOSS
        if i % 2 == 0:
            for cell in cells:
                _shade_cell(cell, "F6F4F0")

    doc.add_page_break()


def _contact_section(doc: Document, ctx: dict) -> None:
    """Advisor contact table. Always rendered (every IA proposal needs
    one); falls back to the document author when override fields are
    empty."""
    ov = ctx.get("overrides") or {}
    advisor_name = (ctx.get("prepared_by") or "").strip()
    advisor_email = ov.get("advisor_email") or ""
    advisor_phone = ov.get("advisor_phone") or ""
    if not (advisor_name or advisor_email or advisor_phone):
        return

    _section_header(doc, ctx, "page.contact", _T(ctx, "page.contact"))
    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "contact.subheading"))
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    table = doc.add_table(rows=2, cols=3)
    headers = [
        _T(ctx, "contact.col_advisor"),
        _T(ctx, "contact.col_email"),
        _T(ctx, "contact.col_phone"),
    ]
    for i, hd in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
        _shade_cell(table.rows[0].cells[i], "010626")
    row_vals = [advisor_name, advisor_email, advisor_phone]
    for i, v in enumerate(row_vals):
        _set_cell_text(table.rows[1].cells[i], v, size_pt=10.5)

    # No trailing page break: Contact is the final section of the document.


def _section_header(doc: Document, ctx: dict, page_tag_key: str, action_title: str) -> None:
    """Per-section title block following Short Brand Guideline §7.1:
    kicker label in Söhne Kräftig + Sunset Ember → big serif title in
    Sometimes Times Medium / Nightblue → thin Nightblue rule.

    Both the kicker and the title use the brand fonts directly so the
    document doesn't rely on Word's "bold" synthesis (which would
    double-weight an already-weighted face).
    """
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(6)
    krun = kicker.add_run(_T(ctx, page_tag_key).upper())
    krun.font.size = Pt(8.5)
    krun.font.color.rgb = SUNSET_EMBER
    krun.font.name = FONT_KRAFTIG
    _apply_font_family(krun.element, FONT_KRAFTIG)
    # Tracked letters per brand "kicker" treatment.
    rpr = krun._r.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "40")   # 40 twentieths-of-a-pt ≈ 0.10em
    rpr.append(spacing)

    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(4)
    hrun = h.add_run(action_title)
    hrun.font.size = Pt(22)
    hrun.font.color.rgb = NIGHTBLUE
    hrun.font.name = FONT_HEADING
    _apply_font_family(hrun.element, FONT_HEADING)

    _add_horizontal_rule(doc, NIGHTBLUE)


def _exec_summary(doc: Document, ctx: dict) -> None:
    _section_header(doc, ctx, "page.exec_summary", ctx.get("exec_title", ""))

    # KPI strip
    kpi_table = doc.add_table(rows=2, cols=3)
    kpi_table.autofit = True
    headers = [
        _T(ctx, "kpi.portfolio_value"),
        _T(ctx, "kpi.defensive_sleeve"),
        _T(ctx, "kpi.risk_profile"),
    ]
    values = [
        f"${ctx.get('portfolio_value', 0):,.0f}",
        f"{ctx.get('defensive_pct', 0):.0f}%",
        profile_label(ctx.get("client", {}).get("profile", ""), ctx.get("lang", I18N_DEFAULT)),
    ]
    subs = [
        _T(ctx, "kpi.reporting_ccy", ccy=ctx.get("client", {}).get("currency", "USD")),
        "USDC · EURC · PAXG",
        _T(ctx, "kpi.positions", n=ctx.get("allocation_count", 0)),
    ]
    for i, (hd, vl, sub) in enumerate(zip(headers, values, subs)):
        _set_cell_text(kpi_table.rows[0].cells[i], hd, bold=True, color=TEXT_MUTED, size_pt=8.5)
        c = kpi_table.rows[1].cells[i]
        c.text = ""
        p1 = c.paragraphs[0]
        r1 = p1.add_run(vl)
        r1.font.size = Pt(20)
        r1.font.bold = True
        r1.font.color.rgb = NIGHTBLUE
        p2 = c.add_paragraph()
        r2 = p2.add_run(sub)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = TEXT_MUTED
        _shade_cell(c, "F6F4F0")
        _shade_cell(kpi_table.rows[0].cells[i], "F6F4F0")

    doc.add_paragraph()

    # "What this means" bullets
    h = doc.add_paragraph()
    hr = h.add_run(_T(ctx, "exhibit.what_this_means"))
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = NIGHTBLUE

    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "exhibit.what_this_means_sub"))
    sr.font.size = Pt(9)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    for b in ctx.get("exec_bullets", []):
        doc.add_paragraph(b, style="List Bullet")

    # Embed the donut exhibit if rasterisation succeeds.
    png = _svg_to_png(ctx.get("donut_svg") or "", width_px=600)
    if png:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(png), width=Cm(10))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(_T(ctx, "exhibit.recommended_alloc_sub"))
        cr.font.size = Pt(8.5)
        cr.font.italic = True
        cr.font.color.rgb = TEXT_MUTED

    doc.add_page_break()


def _allocation_table(doc: Document, ctx: dict) -> None:
    """Per-asset target weights table, tier-bar exhibit and (when the
    client has a recurring-buy plan) the Monthly Buy Schedule.

    No section header of its own: this is the lower half of the merged
    "your portfolio in detail" section and renders straight below
    _portfolio_detail, so the allocation finding sentence appears once."""
    rows = ctx.get("allocation_rows", [])
    h = doc.add_paragraph()
    hr = h.add_run(_T(ctx, "exhibit.per_asset_weights"))
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = NIGHTBLUE

    sub = doc.add_paragraph()
    _sub_key = ("exhibit.per_asset_weights_sub_basket" if ctx.get("is_basket")
                else "exhibit.per_asset_weights_sub")
    sr = sub.add_run(_T(ctx, _sub_key))
    sr.font.size = Pt(9)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = [
        _T(ctx, "table.asset"),
        _T(ctx, "table.tier"),
        _T(ctx, "table.weight_pct"),
        _T(ctx, "table.alloc_usd"),
        _T(ctx, "table.role"),
    ]
    for i, hd in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
        _shade_cell(table.rows[0].cells[i], "010626")

    lang = ctx.get("lang", I18N_DEFAULT)
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        # Asset cell: ticker bold + name
        c = cells[0]
        c.text = ""
        p = c.paragraphs[0]
        r1 = p.add_run(row.get("ticker", ""))
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(f"  {row.get('name', '')}")
        r2.font.size = Pt(9)
        r2.font.color.rgb = TEXT_MUTED
        _set_cell_text(cells[1], tier_label(row.get("tier", ""), lang), size_pt=9.5)
        _set_cell_text(cells[2], f"{row.get('target_pct', 0):.2f}%", align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[3], f"${row.get('target_usd', 0):,.0f}", align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
        _set_cell_text(cells[4], row.get("rationale_tag", ""), size_pt=9.5, color=TEXT_MUTED)
        # Zebra stripe
        if i % 2 == 0:
            for cell in cells:
                _shade_cell(cell, "F6F4F0")

    # Disclosed-omission note when overrides excluded any tickers.
    excluded = (ctx.get("overrides") or {}).get("excluded_tickers") or []
    if excluded:
        note = doc.add_paragraph()
        nr = note.add_run(_T(ctx, "overrides.omitted_note", tickers=", ".join(excluded)))
        nr.font.size = Pt(8.5)
        nr.font.italic = True
        nr.font.color.rgb = TEXT_MUTED

    # Tier-bar exhibit underneath.
    png = _svg_to_png(ctx.get("tier_bar_svg") or "", width_px=720)
    if png:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(png), width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(_T(ctx, "exhibit.alloc_by_tier"))
        cr.font.size = Pt(8.5)
        cr.font.italic = True
        cr.font.color.rgb = TEXT_MUTED

    # Monthly Buy Schedule — rendered here, directly under the
    # recommended allocation, only when the client has a DCA plan.
    _dca_table(doc, ctx)

    doc.add_page_break()


def _overrides_sections(doc: Document, ctx: dict) -> None:
    ov = ctx.get("overrides") or {}
    blocks = [
        ("page.wishes", "overrides.wishes_sub", ov.get("wishes_md")),
        ("page.summary", "overrides.summary_sub", ov.get("summary_md")),
        ("page.execution_plan", "overrides.execution_sub", ov.get("execution_plan_md")),
    ]
    rendered = False
    for tag_key, sub_key, body in blocks:
        if not (body and body.strip()):
            continue
        rendered = True
        _section_header(doc, ctx, tag_key, _T(ctx, tag_key))
        sub = doc.add_paragraph()
        sr = sub.add_run(_T(ctx, sub_key))
        sr.font.size = Pt(9.5)
        sr.font.italic = True
        sr.font.color.rgb = TEXT_MUTED
        _add_md_block(doc, body)
        doc.add_page_break()


def _macro_section(doc: Document, ctx: dict) -> None:
    _section_header(doc, ctx, "page.macro", ctx.get("macro_title", ""))

    h = doc.add_paragraph()
    hr = h.add_run(_T(ctx, "exhibit.regime_read"))
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = NIGHTBLUE

    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "exhibit.regime_read_sub"))
    sr.font.size = Pt(9)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    png = _svg_to_png(ctx.get("regime_gauge_svg") or "", width_px=480)
    if png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(png), width=Cm(8.5))

    body_p = doc.add_paragraph()
    body_p.add_run(ctx.get("macro_paragraph_text", ""))

    inds = ctx.get("top_indicators") or []
    if inds:
        doc.add_paragraph()
        table = doc.add_table(rows=1 + len(inds), cols=3)
        headers = [
            _T(ctx, "table.indicator"),
            _T(ctx, "table.category"),
            _T(ctx, "table.score"),
        ]
        for i, hd in enumerate(headers):
            _set_cell_text(table.rows[0].cells[i], hd, bold=True, color=WHITE, size_pt=9.5)
            _shade_cell(table.rows[0].cells[i], "010626")
        for i, ind in enumerate(inds, start=1):
            cells = table.rows[i].cells
            _set_cell_text(cells[0], ind.get("label", ""), size_pt=9.5)
            _set_cell_text(cells[1], ind.get("category", ""), size_pt=9.5)
            score = ind.get("score")
            _set_cell_text(cells[2],
                           "—" if score is None else f"{score:.0f}",
                           align=WD_ALIGN_PARAGRAPH.RIGHT, size_pt=9.5)
            if i % 2 == 0:
                for cell in cells:
                    _shade_cell(cell, "F6F4F0")

    doc.add_page_break()


def _implementation_section(doc: Document, ctx: dict) -> None:
    lang = ctx.get("lang", I18N_DEFAULT)
    profile = ctx.get("client", {}).get("profile", "")
    heading = t("implementation.heading_tpl", lang,
                profile=profile_label(profile, lang))
    _section_header(doc, ctx, "page.implementation", heading)

    h = doc.add_paragraph()
    hr = h.add_run(_T(ctx, "exhibit.execution_path"))
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = NIGHTBLUE

    sub = doc.add_paragraph()
    sr = sub.add_run(_T(ctx, "exhibit.execution_path_sub"))
    sr.font.size = Pt(9)
    sr.font.italic = True
    sr.font.color.rgb = TEXT_MUTED

    body = doc.add_paragraph()
    body.add_run(ctx.get("implementation_text", ""))

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    hr2 = h2.add_run(_T(ctx, "exhibit.review_cadence"))
    hr2.font.bold = True
    hr2.font.size = Pt(12)
    hr2.font.color.rgb = NIGHTBLUE
    for key in ("review.drift_window", "review.single_drift", "review.rescore_cadence"):
        doc.add_paragraph(_T(ctx, key), style="List Bullet")

    doc.add_page_break()


def _appendix(doc: Document, ctx: dict) -> None:
    lang = ctx.get("lang", I18N_DEFAULT)
    title_en = "Methodology, data sources and important information"
    title_de = "Methodik, Datenquellen und wichtige Hinweise"
    _section_header(doc, ctx, "page.appendix", title_de if lang == "de" else title_en)

    h = doc.add_paragraph()
    hr = h.add_run(_T(ctx, "exhibit.methodology"))
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = NIGHTBLUE
    doc.add_paragraph(_T(ctx, "methodology.body_basket" if ctx.get("is_basket") else "methodology.body"))

    h2 = doc.add_paragraph()
    hr2 = h2.add_run(_T(ctx, "exhibit.data_sources"))
    hr2.font.bold = True
    hr2.font.size = Pt(12)
    hr2.font.color.rgb = NIGHTBLUE
    for key in ("datasources.spot", "datasources.macro", "datasources.onchain", "datasources.history"):
        doc.add_paragraph(_T(ctx, key), style="List Bullet")

    doc.add_paragraph()
    disclaimer = ctx.get("disclaimer") or {}
    h3 = doc.add_paragraph()
    hr3 = h3.add_run(disclaimer.get("title") or "")
    hr3.font.bold = True
    hr3.font.size = Pt(11)
    hr3.font.color.rgb = NIGHTBLUE
    for key in ("body", "tax", "estate"):
        body = disclaimer.get(key)
        if body:
            p = doc.add_paragraph()
            r = p.add_run(body)
            r.font.size = Pt(9.5)

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sr = sign.add_run(_T(ctx, "signoff.prepared_by"))
    sr.font.bold = True
    sr.font.size = Pt(10)
    foot = doc.add_paragraph()
    fr = foot.add_run(f"{ctx.get('prepared_by', '')} · {ctx.get('prepared_date', '')}")
    fr.font.size = Pt(9.5)
    fr.font.color.rgb = TEXT_MUTED

    # Page break so the Contact section (rendered last) starts fresh.
    doc.add_page_break()


# ── Public entry point ──────────────────────────────────────────────


def render_docx(ctx: dict[str, Any]) -> bytes:
    """Build the proposal as a Word document from a proposal context dict.

    Two flavours, switched on ``ctx['proposal_type']``:

    * ``"new"`` (default) — onboarding allocation proposal. Section
      order mirrors Jannick Bröring's WIP IA template: cover → welcome
      → client info → consultation summary → market analysis → merged
      portfolio detail + recommended allocation (with the Monthly Buy
      Schedule when the client has a DCA plan) → strategy → macro
      framing → fazit → wishes → fees → appendix → contact.

    * ``"review"`` — existing-client portfolio review. Inserts three
      review-specific sections after market analysis: current
      holdings + P&L, drift vs target, recommended rebalance trades.
      The target-allocation section is still rendered but reads as
      "what we are rebalancing toward", not "what we propose buying".

    The returned bytes are a complete .docx archive suitable for the
    response body of a FastAPI handler.
    """
    doc = Document()
    theme = _resolve_theme(ctx)
    _set_default_font(doc, theme)
    _page_setup(doc, ctx, theme)

    proposal_type = (ctx.get("proposal_type") or "new").strip().lower()
    is_review = proposal_type == "review"

    _cover(doc, ctx, theme)
    _welcome(doc, ctx)
    _client_info(doc, ctx)
    _consultation_summary(doc, ctx)
    _market_analysis(doc, ctx)
    if is_review:
        _current_holdings(doc, ctx)
        _drift_analysis(doc, ctx)
    else:
        # Per Leonardo (2026-05-19): always show a "current allocation"
        # slot in NEW-client proposals, so advisors can document what
        # the client holds today before proposing the new portfolio.
        _current_allocation(doc, ctx)
    _portfolio_detail(doc, ctx)
    _allocation_table(doc, ctx)
    if is_review:
        _rebalance_actions(doc, ctx)
    _strategy_section(doc, ctx)
    _macro_section(doc, ctx)
    _fazit_section(doc, ctx)
    _wishes_section(doc, ctx)
    _fees_section(doc, ctx)
    # Contact is the closing page, after the appendix — which carries
    # the methodology, data sources and compliance disclaimer.
    _appendix(doc, ctx)
    _contact_section(doc, ctx)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["render_docx"]
